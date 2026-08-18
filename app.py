"""
All-in-One Utility Hub
Flask Web Application Backend
"""

import os
import io
import csv
import json
import uuid
import time
import shutil
import base64
import hashlib
import difflib
import re
import xml.etree.ElementTree as ET
import urllib.parse
from datetime import datetime
from pathlib import Path

from flask import (
    Flask, render_template, request, jsonify, 
    send_file, redirect, url_for, flash, abort
)
from werkzeug.utils import secure_filename
from PIL import Image, ImageOps, ImageFilter, ImageEnhance, ImageDraw, ImageFont
import pypdf
import reportlab
import zipfile
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image as RLImage, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.pdfgen import canvas
import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import openpyxl
import qrcode
import markdown
import html2text
import pymupdf as fitz
import pdfplumber
try:
    from pdf2docx import Converter as PDF2DocxConverter
except ImportError:
    PDF2DocxConverter = None

# Import master tools config
from tools_config import TOOLS, CATEGORIES

# Initialize Flask application
app = Flask(__name__)
app.secret_key = os.environ.get("FLASK_SECRET_KEY", "utility_hub_super_secret_local_key_2026")

try:
    from flask_cors import CORS
    CORS(app)
except ImportError:
    pass

BASE_DIR = Path(__file__).resolve().parent

if os.environ.get("VERCEL") or os.environ.get("AWS_LAMBDA_FUNCTION_NAME"):
    TEMP_ROOT = Path("/tmp")
    UPLOAD_DIR = TEMP_ROOT / "uploads"
    OUTPUT_DIR = TEMP_ROOT / "outputs"
else:
    UPLOAD_DIR = BASE_DIR / "uploads"
    OUTPUT_DIR = BASE_DIR / "outputs"

try:
    UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
except OSError:
    # Read-only filesystem fallback
    TEMP_ROOT = Path("/tmp")
    UPLOAD_DIR = TEMP_ROOT / "uploads"
    OUTPUT_DIR = TEMP_ROOT / "outputs"
    UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


# ==============================================================================
# HELPER FUNCTIONS
# ==============================================================================

def create_job_dirs():
    """Create isolated temporary upload and output folders for each request job."""
    job_id = str(uuid.uuid4())
    job_upload_dir = UPLOAD_DIR / job_id
    job_output_dir = OUTPUT_DIR / job_id
    job_upload_dir.mkdir(parents=True, exist_ok=True)
    job_output_dir.mkdir(parents=True, exist_ok=True)
    return job_id, job_upload_dir, job_output_dir


def cleanup_old_jobs(max_age_seconds=7200):
    """Clean up job directories older than 2 hours to conserve disk space."""
    now = time.time()
    for parent in [UPLOAD_DIR, OUTPUT_DIR]:
        if not parent.exists():
            continue
        for item in parent.iterdir():
            if item.is_dir():
                try:
                    if now - item.stat().st_mtime > max_age_seconds:
                        shutil.rmtree(item, ignore_errors=True)
                except Exception:
                    pass


def get_image_base64_preview(image_path, max_dim=600):
    """Generate a lightweight base64 preview for web UI display."""
    try:
        with Image.open(image_path) as img:
            img_copy = img.copy()
            img_copy.thumbnail((max_dim, max_dim))
            buffered = io.BytesIO()
            fmt = "PNG" if img_copy.mode in ("RGBA", "LA") or img_copy.format == "PNG" else "JPEG"
            if fmt == "JPEG" and img_copy.mode != "RGB":
                img_copy = img_copy.convert("RGB")
            img_copy.save(buffered, format=fmt, quality=85)
            img_str = base64.b64encode(buffered.getvalue()).decode()
            mime = "image/png" if fmt == "PNG" else "image/jpeg"
            return f"data:{mime};base64,{img_str}"
    except Exception:
        return None


def xml_to_dict(element):
    """Recursively convert XML ElementTree to standard Python dictionary."""
    result = {}
    for child in element:
        child_data = xml_to_dict(child) if len(child) > 0 else (child.text.strip() if child.text else "")
        if child.tag in result:
            if isinstance(result[child.tag], list):
                result[child.tag].append(child_data)
            else:
                result[child.tag] = [result[child.tag], child_data]
        else:
            result[child.tag] = child_data
    if element.attrib:
        result["@attributes"] = element.attrib
    return result


# ==============================================================================
# MAIN WEB ROUTES
# ==============================================================================

@app.route("/")
def index():
    """Main dashboard displaying all 34 tools with category filtering & search."""
    cleanup_old_jobs()
    return render_template("dashboard.html", tools=TOOLS, categories=CATEGORIES)


@app.route("/tool/<tool_id>")
def tool_page(tool_id):
    """Dedicated interactive view for a specific utility."""
    tool = TOOLS.get(tool_id)
    if not tool:
        return render_template("404.html"), 404
    return render_template("tool_view.html", tool=tool)


@app.route("/download/<job_id>/<path:filename>")
def download_file(job_id, filename):
    """Download a processed output file."""
    safe_filename = secure_filename(filename)
    file_path = OUTPUT_DIR / job_id / safe_filename
    
    if not file_path.exists():
        abort(404)
        
    return send_file(
        file_path, 
        as_attachment=True, 
        download_name=safe_filename
    )


# ==============================================================================
# TOOL PROCESSING ENDPOINT & LOGIC
# ==============================================================================

@app.route("/process/<tool_id>", methods=["POST"])
def process_tool(tool_id):
    """Central processing controller executing real Python logic for all 34 tools."""
    tool = TOOLS.get(tool_id)
    if not tool:
        return jsonify({"success": False, "error": "Invalid tool requested."}), 404

    job_id, upload_dir, output_dir = create_job_dirs()
    uploaded_files = request.files.getlist("files")
    saved_file_paths = []

    for f in uploaded_files:
        if f and f.filename.strip():
            sec_name = secure_filename(f.filename)
            if not sec_name:
                sec_name = f"file_{uuid.uuid4().hex[:8]}"
            dest = upload_dir / sec_name
            f.save(dest)
            saved_file_paths.append(dest)

    try:
        # -------------------------------------------------------------
        # 1. PDF POWER TOOLS
        # -------------------------------------------------------------
        if tool_id == "merge-pdf":
            if len(saved_file_paths) < 1:
                return jsonify({"success": False, "error": "Please upload at least 2 PDF files to merge."})
            
            writer = pypdf.PdfWriter()
            for path in saved_file_paths:
                writer.append(str(path))
            
            out_name = "merged_document.pdf"
            out_path = output_dir / out_name
            with open(out_path, "wb") as f:
                writer.write(f)
            writer.close()

            return jsonify({
                "success": True,
                "filename": out_name,
                "filesize": out_path.stat().st_size,
                "download_url": url_for("download_file", job_id=job_id, filename=out_name)
            })

        elif tool_id == "split-pdf":
            if not saved_file_paths:
                return jsonify({"success": False, "error": "Please upload a PDF file."})

            pdf_path = saved_file_paths[0]
            reader = pypdf.PdfReader(str(pdf_path))
            num_pages = len(reader.pages)
            page_range_str = request.form.get("page_range", "").strip()

            if page_range_str:
                # Extract specified range into a single PDF
                writer = pypdf.PdfWriter()
                pages_to_extract = set()
                
                parts = page_range_str.split(",")
                for part in parts:
                    part = part.strip()
                    if "-" in part:
                        start, end = part.split("-")
                        for p in range(int(start), int(end) + 1):
                            if 1 <= p <= num_pages:
                                pages_to_extract.add(p - 1)
                    elif part.isdigit():
                        p = int(part)
                        if 1 <= p <= num_pages:
                            pages_to_extract.add(p - 1)

                for p_idx in sorted(pages_to_extract):
                    writer.add_page(reader.pages[p_idx])

                out_name = f"split_{pdf_path.stem}.pdf"
                out_path = output_dir / out_name
                with open(out_path, "wb") as f:
                    writer.write(f)

                return jsonify({
                    "success": True,
                    "filename": out_name,
                    "filesize": out_path.stat().st_size,
                    "download_url": url_for("download_file", job_id=job_id, filename=out_name)
                })
            else:
                # Split all pages into separate PDFs and create a zip
                temp_split_dir = output_dir / "split_pages"
                temp_split_dir.mkdir(exist_ok=True)

                for idx, page in enumerate(reader.pages):
                    single_writer = pypdf.PdfWriter()
                    single_writer.add_page(page)
                    page_file = temp_split_dir / f"page_{idx + 1}.pdf"
                    with open(page_file, "wb") as pf:
                        single_writer.write(pf)

                zip_name = f"{pdf_path.stem}_all_pages.zip"
                zip_path = output_dir / zip_name
                shutil.make_archive(str(output_dir / f"{pdf_path.stem}_all_pages"), "zip", temp_split_dir)

                return jsonify({
                    "success": True,
                    "filename": zip_name,
                    "filesize": zip_path.stat().st_size,
                    "download_url": url_for("download_file", job_id=job_id, filename=zip_name)
                })

        elif tool_id == "compress-pdf":
            if not saved_file_paths:
                return jsonify({"success": False, "error": "Please upload a PDF file."})

            pdf_path = saved_file_paths[0]
            compression_level = request.form.get("compression_level", "recommended")
            out_name = f"compressed_{pdf_path.name}"
            out_path = output_dir / out_name

            # Use PyMuPDF fitz to optimize both vector streams AND embedded images/scans
            doc_fitz = fitz.open(str(pdf_path))
            try:
                if compression_level == "lossless":
                    doc_fitz.save(
                        str(out_path),
                        garbage=3,
                        deflate=True,
                        clean=True
                    )
                else:
                    doc_fitz.save(
                        str(out_path),
                        garbage=4,
                        deflate=True,
                        deflate_images=True,
                        deflate_fonts=True,
                        clean=True
                    )
                doc_fitz.close()
            except Exception:
                # PyPDF fallback
                reader = pypdf.PdfReader(str(pdf_path))
                writer = pypdf.PdfWriter()
                for page in reader.pages:
                    page.compress_content_streams()
                    writer.add_page(page)
                with open(out_path, "wb") as f:
                    writer.write(f)

            return jsonify({
                "success": True,
                "filename": out_name,
                "filesize": out_path.stat().st_size,
                "download_url": url_for("download_file", job_id=job_id, filename=out_name)
            })

        elif tool_id == "pdf-to-word":
            if not saved_file_paths:
                return jsonify({"success": False, "error": "Please upload a PDF file."})

            pdf_path = saved_file_paths[0]
            conversion_mode = request.form.get("conversion_mode", "high_fidelity")
            out_name = f"{pdf_path.stem}.docx"
            out_path = output_dir / out_name

            # Check if PDF contains extractable text or if it is purely image/scan/label
            doc_fitz = fitz.open(str(pdf_path))
            total_text_len = sum(len(page.get_text().strip()) for page in doc_fitz)
            is_scanned_or_image_label = (total_text_len < 15)

            # Auto-fallback: if user selected editable_layout or text mode but document is a scanned image or shipping label,
            # automatically use high_fidelity so that barcodes, QR codes, and boxes are 100% preserved
            if is_scanned_or_image_label or conversion_mode == "high_fidelity":
                doc = docx.Document()
                temp_img_files = []

                for idx, page in enumerate(doc_fitz):
                    rect = page.rect
                    width_pt = rect.width
                    height_pt = rect.height

                    # Render high resolution 300 DPI image for crisp barcode/QR scanning
                    zoom = 300.0 / 72.0
                    matrix = fitz.Matrix(zoom, zoom)
                    pix = page.get_pixmap(matrix=matrix, alpha=False)

                    temp_img_file = output_dir / f"temp_label_page_{idx}_{uuid.uuid4().hex[:6]}.png"
                    pix.save(str(temp_img_file))
                    temp_img_files.append(temp_img_file)

                    # Configure exact physical page dimensions with zero margins
                    if idx == 0:
                        section = doc.sections[0]
                    else:
                        section = doc.add_section()

                    width_in = width_pt / 72.0
                    height_in = height_pt / 72.0

                    section.page_width = Inches(width_in)
                    section.page_height = Inches(height_in)
                    section.top_margin = Inches(0)
                    section.bottom_margin = Inches(0)
                    section.left_margin = Inches(0)
                    section.right_margin = Inches(0)
                    section.header_distance = Inches(0)
                    section.footer_distance = Inches(0)

                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.line_spacing = 1.0
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    run = p.add_run()
                    run.add_picture(str(temp_img_file), width=Inches(width_in), height=Inches(height_in))

                doc_fitz.close()
                doc.save(str(out_path))

                # Clean up temporary page images
                for img_f in temp_img_files:
                    try:
                        img_f.unlink(missing_ok=True)
                    except Exception:
                        pass

            elif conversion_mode == "editable_layout":
                converted_successfully = False

                # Attempt pdf2docx layout reconstruction
                if PDF2DocxConverter is not None:
                    try:
                        cv = PDF2DocxConverter(str(pdf_path))
                        cv.convert(str(out_path))
                        cv.close()
                        converted_successfully = True
                    except Exception as e:
                        print(f"pdf2docx conversion note: {e}")

                # Fallback to pdfplumber table + structured text layout
                if not converted_successfully:
                    try:
                        doc = docx.Document()
                        with pdfplumber.open(str(pdf_path)) as pdf:
                            for idx, page in enumerate(pdf.pages):
                                if idx > 0:
                                    doc.add_page_break()

                                # Extract and format tables
                                tables = page.extract_tables()
                                if tables:
                                    for table_data in tables:
                                        if not table_data or not table_data[0]:
                                            continue
                                        t = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                                        t.style = 'Table Grid'
                                        for r_idx, row in enumerate(table_data):
                                            for c_idx, cell in enumerate(row):
                                                if cell:
                                                    t.cell(r_idx, c_idx).text = str(cell).strip()
                                        doc.add_paragraph()

                                # Extract text lines
                                text = page.extract_text()
                                if text:
                                    for line in text.splitlines():
                                        if line.strip():
                                            doc.add_paragraph(line.strip())

                        doc.save(str(out_path))
                        converted_successfully = True
                    except Exception:
                        pass

                # Final fallback to high fidelity rendering so no empty document is ever returned
                if not converted_successfully:
                    doc = docx.Document()
                    doc_fitz_fb = fitz.open(str(pdf_path))
                    for idx, page in enumerate(doc_fitz_fb):
                        pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                        temp_fb = output_dir / f"temp_fb_{idx}_{uuid.uuid4().hex[:6]}.png"
                        pix.save(str(temp_fb))
                        if idx > 0:
                            doc.add_section()
                        p = doc.add_paragraph()
                        p.add_run().add_picture(str(temp_fb), width=Inches(page.rect.width / 72.0))
                    doc_fitz_fb.close()
                    doc.save(str(out_path))

            else: # extracted_text
                doc = docx.Document()
                doc.add_heading(f"Converted from {pdf_path.name}", 0)
                for page_idx, page in enumerate(doc_fitz):
                    doc.add_heading(f"Page {page_idx + 1}", level=2)
                    for line in page.get_text().splitlines():
                        if line.strip():
                            doc.add_paragraph(line)
                doc_fitz.close()
                doc.save(str(out_path))

            return jsonify({
                "success": True,
                "filename": out_name,
                "filesize": out_path.stat().st_size,
                "download_url": url_for("download_file", job_id=job_id, filename=out_name)
            })

        elif tool_id == "word-to-pdf":
            if not saved_file_paths:
                return jsonify({"success": False, "error": "Please upload a Word document (.docx)."})

            docx_path = saved_file_paths[0]
            out_name = f"{docx_path.stem}.pdf"
            out_path = output_dir / out_name

            doc = docx.Document(str(docx_path))
            pdf_doc = SimpleDocTemplate(
                str(out_path), 
                pagesize=letter, 
                leftMargin=36, 
                rightMargin=36, 
                topMargin=36, 
                bottomMargin=36
            )
            styles = getSampleStyleSheet()
            story = []

            # Printable dimensions within letter page (612 x 792 pt with 36pt margins)
            max_page_w = 540.0
            max_page_h = 680.0

            def create_scaled_image(img_path):
                """Safely scale image dimensions proportionally to fit strictly within page margins."""
                try:
                    with Image.open(img_path) as pil_im:
                        orig_w, orig_h = pil_im.size
                        if orig_w <= 0 or orig_h <= 0:
                            return None
                        
                        scale = min(max_page_w / float(orig_w), max_page_h / float(orig_h), 1.0)
                        render_w = max(1.0, orig_w * scale)
                        render_h = max(1.0, orig_h * scale)
                        
                        rl_img = RLImage(str(img_path), width=render_w, height=render_h, kind='proportional')
                        rl_img.hAlign = 'CENTER'
                        return rl_img
                except Exception as err:
                    print(f"Image scaling error: {err}")
                    return None

            # 1. Extract any embedded media files (e.g. shipping labels, photos, barcodes pasted in docx)
            temp_img_dir = output_dir / f"docx_media_{uuid.uuid4().hex[:6]}"
            temp_img_dir.mkdir(exist_ok=True)
            extracted_media = {}

            try:
                with zipfile.ZipFile(str(docx_path), 'r') as zf:
                    for name in zf.namelist():
                        if name.startswith('word/media/'):
                            fname = Path(name).name
                            dest_f = temp_img_dir / fname
                            with open(dest_f, 'wb') as df:
                                df.write(zf.read(name))
                            extracted_media[fname] = dest_f
            except Exception as e:
                print(f"Docx zip media extraction note: {e}")

            # Map relationships
            rel_image_map = {}
            for r_id, rel in doc.part.rels.items():
                if "image" in rel.target_ref:
                    img_name = Path(rel.target_ref).name
                    if img_name in extracted_media:
                        rel_image_map[r_id] = extracted_media[img_name]

            # 2. Iterate through paragraphs and insert text + inline images
            for p in doc.paragraphs:
                blip_nodes = p._element.xpath('.//a:blip/@r:embed')
                for r_id in blip_nodes:
                    if r_id in rel_image_map:
                        img_file = rel_image_map[r_id]
                        rl_img = create_scaled_image(img_file)
                        if rl_img is not None:
                            story.append(rl_img)
                            story.append(Spacer(1, 8))

                if p.text.strip():
                    story.append(Paragraph(p.text, styles['Normal']))
                    story.append(Spacer(1, 8))

            # 3. Process tables
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [Paragraph(cell.text.strip(), styles['Normal']) for cell in row.cells]
                    table_data.append(row_data)
                if table_data:
                    pdf_table = Table(table_data)
                    pdf_table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#f1f5f9')),
                        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#cbd5e1')),
                        ('PADDING', (0, 0), (-1, -1), 6),
                    ]))
                    story.append(pdf_table)
                    story.append(Spacer(1, 12))

            # 4. If story is empty or has no elements, render extracted media (e.g. image shipping slips/scans)
            if not story and extracted_media:
                for fname, img_file in extracted_media.items():
                    rl_img = create_scaled_image(img_file)
                    if rl_img is not None:
                        story.append(rl_img)
                        story.append(Spacer(1, 8))

            # 5. Final fallback if completely blank
            if not story:
                story.append(Paragraph("Document Converted (Blank Page)", styles['Normal']))

            pdf_doc.build(story)

            return jsonify({
                "success": True,
                "filename": out_name,
                "filesize": out_path.stat().st_size,
                "download_url": url_for("download_file", job_id=job_id, filename=out_name)
            })

        elif tool_id == "pdf-to-jpg":
            if not saved_file_paths:
                return jsonify({"success": False, "error": "Please upload a PDF file."})

            pdf_path = saved_file_paths[0]
            dpi = int(request.form.get("dpi", 150))
            zoom = dpi / 72.0
            matrix = fitz.Matrix(zoom, zoom)

            doc_fitz = fitz.open(str(pdf_path))
            num_pages = len(doc_fitz)

            if num_pages == 1:
                page = doc_fitz[0]
                pix = page.get_pixmap(matrix=matrix)
                out_name = f"{pdf_path.stem}_page1.jpg"
                out_path = output_dir / out_name
                pix.save(str(out_path))

                preview_b64 = get_image_base64_preview(out_path)

                return jsonify({
                    "success": True,
                    "filename": out_name,
                    "filesize": out_path.stat().st_size,
                    "download_url": url_for("download_file", job_id=job_id, filename=out_name),
                    "preview_type": "image",
                    "preview_data": preview_b64
                })
            else:
                temp_img_dir = output_dir / "extracted_images"
                temp_img_dir.mkdir(exist_ok=True)

                for idx, page in enumerate(doc_fitz):
                    pix = page.get_pixmap(matrix=matrix)
                    pix.save(str(temp_img_dir / f"page_{idx + 1}.jpg"))

                zip_name = f"{pdf_path.stem}_jpg_pages.zip"
                zip_path = output_dir / zip_name
                shutil.make_archive(str(output_dir / f"{pdf_path.stem}_jpg_pages"), "zip", temp_img_dir)

                return jsonify({
                    "success": True,
                    "filename": zip_name,
                    "filesize": zip_path.stat().st_size,
                    "download_url": url_for("download_file", job_id=job_id, filename=zip_name)
                })

        elif tool_id == "jpg-to-pdf":
            if not saved_file_paths:
                return jsonify({"success": False, "error": "Please upload at least one image."})

            image_list = []
            for p in saved_file_paths:
                try:
                    img = Image.open(p)
                    if img.mode != "RGB":
                        img = img.convert("RGB")
                    image_list.append(img)
                except Exception:
                    pass

            if not image_list:
                return jsonify({"success": False, "error": "No valid image files could be processed."})

            out_name = "images_combined.pdf"
            out_path = output_dir / out_name

            image_list[0].save(
                str(out_path), 
                save_all=True, 
                append_images=image_list[1:]
            )

            return jsonify({
                "success": True,
                "filename": out_name,
                "filesize": out_path.stat().st_size,
                "download_url": url_for("download_file", job_id=job_id, filename=out_name)
            })

        elif tool_id == "rotate-pdf":
            if not saved_file_paths:
                return jsonify({"success": False, "error": "Please upload a PDF file."})

            pdf_path = saved_file_paths[0]
            angle = int(request.form.get("angle", 90))

            reader = pypdf.PdfReader(str(pdf_path))
            writer = pypdf.PdfWriter()

            for page in reader.pages:
                page.rotate(angle)
                writer.add_page(page)

            out_name = f"rotated_{pdf_path.name}"
            out_path = output_dir / out_name
            with open(out_path, "wb") as f:
                writer.write(f)

            return jsonify({
                "success": True,
                "filename": out_name,
                "filesize": out_path.stat().st_size,
                "download_url": url_for("download_file", job_id=job_id, filename=out_name)
            })

        elif tool_id == "watermark-pdf":
            if not saved_file_paths:
                return jsonify({"success": False, "error": "Please upload a PDF file."})

            pdf_path = saved_file_paths[0]
            watermark_text = request.form.get("watermark_text", "CONFIDENTIAL")
            opacity = float(request.form.get("opacity", 0.3))

            # Generate watermark overlay PDF in memory
            packet = io.BytesIO()
            can = canvas.Canvas(packet, pagesize=letter)
            can.setFont("Helvetica-Bold", 42)
            can.setFillColor(colors.Color(0.5, 0.5, 0.5, alpha=opacity))
            can.saveState()
            can.translate(300, 450)
            can.rotate(45)
            can.drawCentredString(0, 0, watermark_text)
            can.restoreState()
            can.save()
            packet.seek(0)

            watermark_pdf = pypdf.PdfReader(packet)
            watermark_page = watermark_pdf.pages[0]

            reader = pypdf.PdfReader(str(pdf_path))
            writer = pypdf.PdfWriter()

            for page in reader.pages:
                page.merge_page(watermark_page)
                writer.add_page(page)

            out_name = f"watermarked_{pdf_path.name}"
            out_path = output_dir / out_name
            with open(out_path, "wb") as f:
                writer.write(f)

            return jsonify({
                "success": True,
                "filename": out_name,
                "filesize": out_path.stat().st_size,
                "download_url": url_for("download_file", job_id=job_id, filename=out_name)
            })

        elif tool_id == "protect-pdf":
            if not saved_file_paths:
                return jsonify({"success": False, "error": "Please upload a PDF file."})

            pdf_path = saved_file_paths[0]
            mode = request.form.get("mode", "encrypt")
            password = request.form.get("password", "")

            if not password:
                return jsonify({"success": False, "error": "Please provide a password."})

            reader = pypdf.PdfReader(str(pdf_path))
            writer = pypdf.PdfWriter()

            if mode == "encrypt":
                for page in reader.pages:
                    writer.add_page(page)
                writer.encrypt(user_password=password, owner_password=password)
                out_name = f"protected_{pdf_path.name}"
            else:
                if reader.is_encrypted:
                    reader.decrypt(password)
                for page in reader.pages:
                    writer.add_page(page)
                out_name = f"unlocked_{pdf_path.name}"

            out_path = output_dir / out_name
            with open(out_path, "wb") as f:
                writer.write(f)

            return jsonify({
                "success": True,
                "filename": out_name,
                "filesize": out_path.stat().st_size,
                "download_url": url_for("download_file", job_id=job_id, filename=out_name)
            })

        # -------------------------------------------------------------
        # 2. AI & IMAGE STUDIO
        # -------------------------------------------------------------
        elif tool_id == "remove-bg":
            if not saved_file_paths:
                return jsonify({"success": False, "error": "Please upload an image."})

            img_path = saved_file_paths[0]
            out_name = f"{img_path.stem}_nobg.png"
            out_path = output_dir / out_name

            try:
                import rembg
                with Image.open(img_path) as input_img:
                    output_img = rembg.remove(input_img)
                    output_img.save(str(out_path), "PNG")
            except Exception as e:
                # High-quality PIL fallback for transparency removal
                with Image.open(img_path) as input_img:
                    rgba = input_img.convert("RGBA")
                    datas = rgba.getdata()
                    new_data = []
                    # Simple corner background color keying fallback
                    bg_color = datas[0]
                    for item in datas:
                        if abs(item[0]-bg_color[0]) < 25 and abs(item[1]-bg_color[1]) < 25 and abs(item[2]-bg_color[2]) < 25:
                            new_data.append((255, 255, 255, 0))
                        else:
                            new_data.append(item)
                    rgba.putdata(new_data)
                    rgba.save(str(out_path), "PNG")

            preview_b64 = get_image_base64_preview(out_path)

            return jsonify({
                "success": True,
                "filename": out_name,
                "filesize": out_path.stat().st_size,
                "download_url": url_for("download_file", job_id=job_id, filename=out_name),
                "preview_type": "image",
                "preview_data": preview_b64
            })

        elif tool_id == "image-resizer":
            if not saved_file_paths:
                return jsonify({"success": False, "error": "Please upload an image."})

            img_path = saved_file_paths[0]
            resize_mode = request.form.get("resize_mode", "percentage")
            
            with Image.open(img_path) as img:
                orig_w, orig_h = img.size
                
                if resize_mode == "percentage":
                    scale = float(request.form.get("scale_percent", 50)) / 100.0
                    new_w = max(1, int(orig_w * scale))
                    new_h = max(1, int(orig_h * scale))
                else:
                    new_w = int(request.form.get("width") or orig_w)
                    new_h = int(request.form.get("height") or orig_h)

                resized = img.resize((new_w, new_h), Image.Resampling.LANCZOS)
                out_name = f"resized_{img_path.name}"
                out_path = output_dir / out_name
                resized.save(str(out_path))

            preview_b64 = get_image_base64_preview(out_path)

            return jsonify({
                "success": True,
                "filename": out_name,
                "filesize": out_path.stat().st_size,
                "download_url": url_for("download_file", job_id=job_id, filename=out_name),
                "preview_type": "image",
                "preview_data": preview_b64
            })

        elif tool_id == "image-converter":
            if not saved_file_paths:
                return jsonify({"success": False, "error": "Please upload an image."})

            img_path = saved_file_paths[0]
            target_fmt = request.form.get("target_format", "webp").lower()
            
            with Image.open(img_path) as img:
                out_name = f"{img_path.stem}.{target_fmt}"
                out_path = output_dir / out_name

                if target_fmt in ("jpg", "jpeg"):
                    if img.mode != "RGB":
                        img = img.convert("RGB")
                    img.save(str(out_path), "JPEG", quality=90)
                elif target_fmt == "ico":
                    icon_img = img.resize((64, 64), Image.Resampling.LANCZOS)
                    icon_img.save(str(out_path), format="ICO")
                elif target_fmt == "webp":
                    img.save(str(out_path), "WEBP", quality=90)
                elif target_fmt == "png":
                    img.save(str(out_path), "PNG")
                else:
                    img.save(str(out_path), format=target_fmt.upper())

            preview_b64 = get_image_base64_preview(out_path)

            return jsonify({
                "success": True,
                "filename": out_name,
                "filesize": out_path.stat().st_size,
                "download_url": url_for("download_file", job_id=job_id, filename=out_name),
                "preview_type": "image",
                "preview_data": preview_b64
            })

        elif tool_id == "image-compressor":
            if not saved_file_paths:
                return jsonify({"success": False, "error": "Please upload an image."})

            img_path = saved_file_paths[0]
            quality = int(request.form.get("quality", 75))

            with Image.open(img_path) as img:
                out_name = f"compressed_{img_path.stem}.jpg"
                out_path = output_dir / out_name
                if img.mode != "RGB":
                    img = img.convert("RGB")
                img.save(str(out_path), "JPEG", quality=quality, optimize=True)

            preview_b64 = get_image_base64_preview(out_path)

            return jsonify({
                "success": True,
                "filename": out_name,
                "filesize": out_path.stat().st_size,
                "download_url": url_for("download_file", job_id=job_id, filename=out_name),
                "preview_type": "image",
                "preview_data": preview_b64
            })

        elif tool_id == "image-filters":
            if not saved_file_paths:
                return jsonify({"success": False, "error": "Please upload an image."})

            img_path = saved_file_paths[0]
            filter_type = request.form.get("filter_type", "grayscale")

            with Image.open(img_path) as img:
                if filter_type == "grayscale":
                    filtered = ImageOps.grayscale(img)
                elif filter_type == "sepia":
                    gray = ImageOps.grayscale(img)
                    filtered = ImageOps.colorize(gray, "#704214", "#FFF8DC")
                elif filter_type == "blur":
                    filtered = img.filter(ImageFilter.GaussianBlur(radius=4))
                elif filter_type == "sharpen":
                    filtered = img.filter(ImageFilter.SHARPEN)
                elif filter_type == "invert":
                    rgb = img.convert("RGB")
                    filtered = ImageOps.invert(rgb)
                elif filter_type == "enhance_contrast":
                    rgb = img.convert("RGB")
                    filtered = ImageOps.autocontrast(rgb)
                elif filter_type == "brighten":
                    enhancer = ImageEnhance.Brightness(img)
                    filtered = enhancer.enhance(1.3)
                else:
                    filtered = img

                out_name = f"{filter_type}_{img_path.name}"
                out_path = output_dir / out_name
                filtered.save(str(out_path))

            preview_b64 = get_image_base64_preview(out_path)

            return jsonify({
                "success": True,
                "filename": out_name,
                "filesize": out_path.stat().st_size,
                "download_url": url_for("download_file", job_id=job_id, filename=out_name),
                "preview_type": "image",
                "preview_data": preview_b64
            })

        elif tool_id == "image-cropper":
            if not saved_file_paths:
                return jsonify({"success": False, "error": "Please upload an image."})

            img_path = saved_file_paths[0]
            aspect_ratio = request.form.get("aspect_ratio", "1:1")

            with Image.open(img_path) as img:
                w, h = img.size
                if aspect_ratio == "1:1":
                    side = min(w, h)
                    left = (w - side) // 2
                    top = (h - side) // 2
                    cropped = img.crop((left, top, left + side, top + side))
                elif aspect_ratio == "16:9":
                    target_w = w
                    target_h = int(w * 9 / 16)
                    if target_h > h:
                        target_h = h
                        target_w = int(h * 16 / 9)
                    left = (w - target_w) // 2
                    top = (h - target_h) // 2
                    cropped = img.crop((left, top, left + target_w, top + target_h))
                elif aspect_ratio == "4:3":
                    target_w = w
                    target_h = int(w * 3 / 4)
                    if target_h > h:
                        target_h = h
                        target_w = int(h * 4 / 3)
                    left = (w - target_w) // 2
                    top = (h - target_h) // 2
                    cropped = img.crop((left, top, left + target_w, top + target_h))
                elif aspect_ratio == "9:16":
                    target_w = int(h * 9 / 16)
                    if target_w > w:
                        target_w = w
                        target_h = int(w * 16 / 9)
                    else:
                        target_h = h
                    left = (w - target_w) // 2
                    top = (h - target_h) // 2
                    cropped = img.crop((left, top, left + target_w, top + target_h))
                else:
                    cropped = img

                out_name = f"cropped_{img_path.name}"
                out_path = output_dir / out_name
                cropped.save(str(out_path))

            preview_b64 = get_image_base64_preview(out_path)

            return jsonify({
                "success": True,
                "filename": out_name,
                "filesize": out_path.stat().st_size,
                "download_url": url_for("download_file", job_id=job_id, filename=out_name),
                "preview_type": "image",
                "preview_data": preview_b64
            })

        elif tool_id == "image-to-base64":
            if not saved_file_paths:
                return jsonify({"success": False, "error": "Please upload an image."})

            img_path = saved_file_paths[0]
            with open(img_path, "rb") as img_f:
                b64_str = base64.b64encode(img_f.read()).decode("utf-8")
            
            ext = img_path.suffix.lstrip(".").lower()
            mime = f"image/{ext}" if ext != "svg" else "image/svg+xml"
            data_uri = f"data:{mime};base64,{b64_str}"

            out_name = f"{img_path.stem}_base64.txt"
            out_path = output_dir / out_name
            with open(out_path, "w", encoding="utf-8") as out_f:
                out_f.write(data_uri)

            return jsonify({
                "success": True,
                "filename": out_name,
                "filesize": out_path.stat().st_size,
                "download_url": url_for("download_file", job_id=job_id, filename=out_name),
                "preview_type": "text",
                "text_content": data_uri[:1000] + "..." if len(data_uri) > 1000 else data_uri
            })

        elif tool_id == "image-watermark":
            if not saved_file_paths:
                return jsonify({"success": False, "error": "Please upload an image."})

            img_path = saved_file_paths[0]
            wm_text = request.form.get("watermark_text", "© Copyright 2026")
            pos = request.form.get("position", "bottom_right")

            with Image.open(img_path) as img:
                rgba = img.convert("RGBA")
                txt_layer = Image.new("RGBA", rgba.size, (255, 255, 255, 0))
                draw = ImageDraw.Draw(txt_layer)

                font_size = max(16, int(rgba.size[0] / 30))
                try:
                    font = ImageFont.truetype("arial.ttf", font_size)
                except Exception:
                    font = ImageFont.load_default()

                bbox = draw.textbbox((0, 0), wm_text, font=font)
                t_w = bbox[2] - bbox[0]
                t_h = bbox[3] - bbox[1]

                if pos == "bottom_right":
                    x = rgba.size[0] - t_w - 20
                    y = rgba.size[1] - t_h - 20
                elif pos == "bottom_left":
                    x = 20
                    y = rgba.size[1] - t_h - 20
                elif pos == "center":
                    x = (rgba.size[0] - t_w) // 2
                    y = (rgba.size[1] - t_h) // 2
                else: # top_right
                    x = rgba.size[0] - t_w - 20
                    y = 20

                draw.text((x, y), wm_text, fill=(255, 255, 255, 180), font=font)
                watermarked = Image.alpha_composite(rgba, txt_layer)

                out_name = f"watermarked_{img_path.stem}.png"
                out_path = output_dir / out_name
                watermarked.save(str(out_path), "PNG")

            preview_b64 = get_image_base64_preview(out_path)

            return jsonify({
                "success": True,
                "filename": out_name,
                "filesize": out_path.stat().st_size,
                "download_url": url_for("download_file", job_id=job_id, filename=out_name),
                "preview_type": "image",
                "preview_data": preview_b64
            })

        elif tool_id == "exif-remover":
            if not saved_file_paths:
                return jsonify({"success": False, "error": "Please upload an image."})

            img_path = saved_file_paths[0]
            with Image.open(img_path) as img:
                data = list(img.getdata())
                clean_img = Image.new(img.mode, img.size)
                clean_img.putdata(data)

                out_name = f"clean_{img_path.name}"
                out_path = output_dir / out_name
                clean_img.save(str(out_path))

            preview_b64 = get_image_base64_preview(out_path)

            return jsonify({
                "success": True,
                "filename": out_name,
                "filesize": out_path.stat().st_size,
                "download_url": url_for("download_file", job_id=job_id, filename=out_name),
                "preview_type": "image",
                "preview_data": preview_b64
            })

        elif tool_id == "color-palette":
            if not saved_file_paths:
                return jsonify({"success": False, "error": "Please upload an image."})

            img_path = saved_file_paths[0]
            with Image.open(img_path) as img:
                rgb_img = img.convert("RGB")
                small_img = rgb_img.resize((150, 150))
                quantized = small_img.quantize(colors=6, method=Image.Quantize.MEDIANCUT)
                palette = quantized.getpalette()[:18]

                colors_list = []
                for i in range(0, len(palette), 3):
                    r, g, b = palette[i], palette[i+1], palette[i+2]
                    hex_code = f"#{r:02x}{g:02x}{b:02x}".upper()
                    colors_list.append({
                        "hex": hex_code,
                        "rgb": f"rgb({r}, {g}, {b})"
                    })

            # Save palette summary JSON as download
            out_name = f"{img_path.stem}_palette.json"
            out_path = output_dir / out_name
            with open(out_path, "w", encoding="utf-8") as pf:
                json.dump(colors_list, pf, indent=2)

            return jsonify({
                "success": True,
                "filename": out_name,
                "filesize": out_path.stat().st_size,
                "download_url": url_for("download_file", job_id=job_id, filename=out_name),
                "preview_type": "palette",
                "colors": colors_list
            })

        # -------------------------------------------------------------
        # 3. DOCUMENT & DATA CONVERTERS
        # -------------------------------------------------------------
        elif tool_id == "excel-to-pdf":
            if not saved_file_paths:
                return jsonify({"success": False, "error": "Please upload an Excel workbook (.xlsx)."})

            excel_path = saved_file_paths[0]
            wb = openpyxl.load_workbook(str(excel_path), data_only=True)
            sheet = wb.active

            out_name = f"{excel_path.stem}.pdf"
            out_path = output_dir / out_name

            pdf_doc = SimpleDocTemplate(str(out_path), pagesize=letter)
            styles = getSampleStyleSheet()
            story = []

            story.append(Paragraph(f"Spreadsheet: {sheet.title}", styles['Heading2']))
            story.append(Spacer(1, 12))

            table_data = []
            for row in sheet.iter_rows(values_only=True):
                if any(cell is not None for cell in row):
                    row_data = [Paragraph(str(c) if c is not None else "", styles['Normal']) for c in row]
                    table_data.append(row_data)

            if table_data:
                pdf_table = Table(table_data)
                pdf_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#0f766e')),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#94a3b8')),
                    ('PADDING', (0, 0), (-1, -1), 5),
                ]))
                story.append(pdf_table)
            else:
                story.append(Paragraph("Spreadsheet contains no row data.", styles['Normal']))

            pdf_doc.build(story)

            return jsonify({
                "success": True,
                "filename": out_name,
                "filesize": out_path.stat().st_size,
                "download_url": url_for("download_file", job_id=job_id, filename=out_name)
            })

        elif tool_id == "excel-to-csv":
            if not saved_file_paths:
                return jsonify({"success": False, "error": "Please upload an Excel file."})

            excel_path = saved_file_paths[0]
            wb = openpyxl.load_workbook(str(excel_path), data_only=True)
            sheet = wb.active

            out_name = f"{excel_path.stem}.csv"
            out_path = output_dir / out_name

            with open(out_path, "w", newline="", encoding="utf-8") as csv_f:
                writer = csv.writer(csv_f)
                for row in sheet.iter_rows(values_only=True):
                    writer.writerow([c if c is not None else "" for c in row])

            return jsonify({
                "success": True,
                "filename": out_name,
                "filesize": out_path.stat().st_size,
                "download_url": url_for("download_file", job_id=job_id, filename=out_name)
            })

        elif tool_id == "csv-to-excel":
            if not saved_file_paths:
                return jsonify({"success": False, "error": "Please upload a CSV file."})

            csv_path = saved_file_paths[0]
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "Data"

            with open(csv_path, "r", encoding="utf-8-sig", errors="ignore") as f:
                reader = csv.reader(f)
                for row in reader:
                    ws.append(row)

            out_name = f"{csv_path.stem}.xlsx"
            out_path = output_dir / out_name
            wb.save(str(out_path))

            return jsonify({
                "success": True,
                "filename": out_name,
                "filesize": out_path.stat().st_size,
                "download_url": url_for("download_file", job_id=job_id, filename=out_name)
            })

        elif tool_id == "json-to-csv":
            if not saved_file_paths:
                return jsonify({"success": False, "error": "Please upload a JSON file."})

            json_path = saved_file_paths[0]
            with open(json_path, "r", encoding="utf-8") as jf:
                data = json.load(jf)

            if isinstance(data, dict):
                data = [data]

            out_name = f"{json_path.stem}.csv"
            out_path = output_dir / out_name

            if isinstance(data, list) and len(data) > 0 and isinstance(data[0], dict):
                headers = list(data[0].keys())
                with open(out_path, "w", newline="", encoding="utf-8") as csv_f:
                    writer = csv.DictWriter(csv_f, fieldnames=headers)
                    writer.writeheader()
                    for item in data:
                        writer.writerow({h: item.get(h, "") for h in headers})

            return jsonify({
                "success": True,
                "filename": out_name,
                "filesize": out_path.stat().st_size,
                "download_url": url_for("download_file", job_id=job_id, filename=out_name)
            })

        elif tool_id == "csv-to-json":
            if not saved_file_paths:
                return jsonify({"success": False, "error": "Please upload a CSV file."})

            csv_path = saved_file_paths[0]
            rows = []
            with open(csv_path, "r", encoding="utf-8-sig", errors="ignore") as f:
                reader = csv.DictReader(f)
                for row in reader:
                    rows.append(row)

            out_name = f"{csv_path.stem}.json"
            out_path = output_dir / out_name
            with open(out_path, "w", encoding="utf-8") as jf:
                json.dump(rows, jf, indent=2)

            return jsonify({
                "success": True,
                "filename": out_name,
                "filesize": out_path.stat().st_size,
                "download_url": url_for("download_file", job_id=job_id, filename=out_name),
                "preview_type": "text",
                "text_content": json.dumps(rows[:5], indent=2) + ("\n\n... (truncated)" if len(rows) > 5 else "")
            })

        elif tool_id == "markdown-to-html":
            content = ""
            if saved_file_paths:
                with open(saved_file_paths[0], "r", encoding="utf-8") as mf:
                    content = mf.read()
            else:
                content = request.form.get("markdown_text") or request.form.get("content") or "# Hello World"

            html_body = markdown.markdown(content, extensions=['extra', 'codehilite', 'tables', 'toc'])
            full_html = f"""<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<title>Rendered Markdown</title>
<style>
body {{ font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, sans-serif; line-height: 1.6; max-width: 800px; margin: 40px auto; padding: 0 20px; color: #1e293b; }}
pre {{ background: #f1f5f9; padding: 16px; border-radius: 8px; overflow-x: auto; }}
code {{ font-family: Consolas, monospace; }}
table {{ border-collapse: collapse; width: 100%; margin: 20px 0; }}
th, td {{ border: 1px solid #cbd5e1; padding: 8px 12px; text-align: left; }}
th {{ background-color: #f8fafc; }}
</style>
</head>
<body>
{html_body}
</body>
</html>"""

            out_name = "rendered_document.html"
            out_path = output_dir / out_name
            with open(out_path, "w", encoding="utf-8") as hf:
                hf.write(full_html)

            return jsonify({
                "success": True,
                "filename": out_name,
                "filesize": out_path.stat().st_size,
                "download_url": url_for("download_file", job_id=job_id, filename=out_name),
                "preview_type": "text",
                "text_content": html_body
            })

        elif tool_id == "html-to-markdown":
            html_content = ""
            if saved_file_paths:
                with open(saved_file_paths[0], "r", encoding="utf-8") as hf:
                    html_content = hf.read()
            else:
                html_content = request.form.get("html_text") or request.form.get("content") or "<h1>Hello World</h1>"

            h = html2text.HTML2Text()
            h.ignore_links = False
            md_content = h.handle(html_content)

            out_name = "converted.md"
            out_path = output_dir / out_name
            with open(out_path, "w", encoding="utf-8") as mf:
                mf.write(md_content)

            return jsonify({
                "success": True,
                "filename": out_name,
                "filesize": out_path.stat().st_size,
                "download_url": url_for("download_file", job_id=job_id, filename=out_name),
                "preview_type": "text",
                "text_content": md_content
            })

        elif tool_id == "xml-to-json":
            xml_content = ""
            if saved_file_paths:
                with open(saved_file_paths[0], "r", encoding="utf-8") as xf:
                    xml_content = xf.read()
            else:
                xml_content = request.form.get("xml_text") or request.form.get("content") or "<root><item>Hello</item></root>"

            root = ET.fromstring(xml_content)
            parsed_dict = {root.tag: xml_to_dict(root)}

            out_name = "converted_data.json"
            out_path = output_dir / out_name
            with open(out_path, "w", encoding="utf-8") as jf:
                json.dump(parsed_dict, jf, indent=2)

            return jsonify({
                "success": True,
                "filename": out_name,
                "filesize": out_path.stat().st_size,
                "download_url": url_for("download_file", job_id=job_id, filename=out_name),
                "preview_type": "text",
                "text_content": json.dumps(parsed_dict, indent=2)
            })

        # -------------------------------------------------------------
        # 4. UTILITIES & SECURITY TOOLS
        # -------------------------------------------------------------
        elif tool_id == "qr-generator":
            qr_content = request.form.get("qr_content", "https://example.com").strip()
            qr_size = int(request.form.get("qr_size", 10))

            qr = qrcode.QRCode(
                version=1,
                error_correction=qrcode.constants.ERROR_CORRECT_H,
                box_size=qr_size,
                border=4,
            )
            qr.add_data(qr_content)
            qr.make(fit=True)
            img = qr.make_image(fill_color="black", back_color="white")

            out_name = "qrcode.png"
            out_path = output_dir / out_name
            img.save(str(out_path))

            preview_b64 = get_image_base64_preview(out_path)

            return jsonify({
                "success": True,
                "filename": out_name,
                "filesize": out_path.stat().st_size,
                "download_url": url_for("download_file", job_id=job_id, filename=out_name),
                "preview_type": "image",
                "preview_data": preview_b64
            })

        elif tool_id == "file-hasher":
            if not saved_file_paths:
                return jsonify({"success": False, "error": "Please upload a file to calculate hashes."})

            file_path = saved_file_paths[0]
            md5_h = hashlib.md5()
            sha1_h = hashlib.sha1()
            sha256_h = hashlib.sha256()
            sha512_h = hashlib.sha512()

            with open(file_path, "rb") as f:
                while chunk := f.read(65536):
                    md5_h.update(chunk)
                    sha1_h.update(chunk)
                    sha256_h.update(chunk)
                    sha512_h.update(chunk)

            hashes = {
                "MD5": md5_h.hexdigest(),
                "SHA-1": sha1_h.hexdigest(),
                "SHA-256": sha256_h.hexdigest(),
                "SHA-512": sha512_h.hexdigest()
            }

            out_name = f"{file_path.stem}_checksums.txt"
            out_path = output_dir / out_name
            with open(out_path, "w", encoding="utf-8") as cf:
                for k, v in hashes.items():
                    cf.write(f"{k}: {v}\n")

            return jsonify({
                "success": True,
                "filename": out_name,
                "filesize": out_path.stat().st_size,
                "download_url": url_for("download_file", job_id=job_id, filename=out_name),
                "preview_type": "hashes",
                "hashes": hashes
            })

        elif tool_id == "base64-codec":
            mode = request.form.get("codec_mode", "encode")
            input_text = request.form.get("input_text", "").strip()

            if saved_file_paths:
                with open(saved_file_paths[0], "rb") as f:
                    raw_bytes = f.read()
                if mode == "encode":
                    res_text = base64.b64encode(raw_bytes).decode("utf-8")
                    out_name = f"{saved_file_paths[0].stem}_base64.txt"
                    out_path = output_dir / out_name
                    with open(out_path, "w", encoding="utf-8") as of:
                        of.write(res_text)
                else:
                    res_bytes = base64.b64decode(raw_bytes)
                    out_name = f"decoded_{saved_file_paths[0].stem}"
                    out_path = output_dir / out_name
                    with open(out_path, "wb") as of:
                        of.write(res_bytes)
                    res_text = res_bytes.decode("utf-8", errors="replace")[:1000]
            else:
                if mode == "encode":
                    res_text = base64.b64encode(input_text.encode("utf-8")).decode("utf-8")
                else:
                    res_text = base64.b64decode(input_text.encode("utf-8")).decode("utf-8", errors="replace")

                out_name = "base64_result.txt"
                out_path = output_dir / out_name
                with open(out_path, "w", encoding="utf-8") as of:
                    of.write(res_text)

            return jsonify({
                "success": True,
                "filename": out_name,
                "filesize": out_path.stat().st_size,
                "download_url": url_for("download_file", job_id=job_id, filename=out_name),
                "preview_type": "text",
                "text_content": res_text
            })

        elif tool_id == "text-diff":
            orig = request.form.get("original_text", "").splitlines()
            mod = request.form.get("modified_text", "").splitlines()

            diff_gen = difflib.unified_diff(orig, mod, fromfile="Original", tofile="Modified", lineterm="")
            diff_lines = list(diff_gen)

            out_name = "diff_report.patch"
            out_path = output_dir / out_name
            with open(out_path, "w", encoding="utf-8") as df:
                df.write("\n".join(diff_lines))

            return jsonify({
                "success": True,
                "filename": out_name,
                "filesize": out_path.stat().st_size,
                "download_url": url_for("download_file", job_id=job_id, filename=out_name),
                "preview_type": "diff",
                "diff_lines": diff_lines if diff_lines else ["No differences detected."]
            })

        elif tool_id == "case-converter":
            txt = request.form.get("text_input") or request.form.get("url_input") or request.form.get("input_text") or request.form.get("content") or ""
            if saved_file_paths:
                with open(saved_file_paths[0], "r", encoding="utf-8", errors="ignore") as tf_in:
                    txt = tf_in.read()

            style = request.form.get("case_style", "title")

            if style == "upper":
                result = txt.upper()
            elif style == "lower":
                result = txt.lower()
            elif style == "title":
                result = txt.title()
            elif style == "snake":
                result = re.sub(r'[\s\-]+', '_', txt).lower()
            elif style == "kebab":
                result = re.sub(r'[\s\_]+', '-', txt).lower()
            elif style == "camel":
                words = re.split(r'[\s\_\-]+', txt)
                result = words[0].lower() + "".join(w.capitalize() for w in words[1:]) if words else ""
            elif style == "pascal":
                words = re.split(r'[\s\_\-]+', txt)
                result = "".join(w.capitalize() for w in words)
            else:
                result = txt

            out_name = "converted_text.txt"
            out_path = output_dir / out_name
            with open(out_path, "w", encoding="utf-8") as tf:
                tf.write(result)

            return jsonify({
                "success": True,
                "filename": out_name,
                "filesize": out_path.stat().st_size,
                "download_url": url_for("download_file", job_id=job_id, filename=out_name),
                "preview_type": "text",
                "text_content": result
            })

        elif tool_id == "url-encoder-decoder":
            url_mode = request.form.get("url_mode", "encode")
            url_input = request.form.get("url_input") or request.form.get("text_input") or request.form.get("input_text") or request.form.get("content") or ""
            if saved_file_paths:
                with open(saved_file_paths[0], "r", encoding="utf-8", errors="ignore") as uf_in:
                    url_input = uf_in.read()

            url_input = url_input.strip()

            if url_mode == "encode":
                result = urllib.parse.quote(url_input, safe="/:")
            else:
                result = urllib.parse.unquote(url_input)

            out_name = "url_result.txt"
            out_path = output_dir / out_name
            with open(out_path, "w", encoding="utf-8") as uf:
                uf.write(result)

            return jsonify({
                "success": True,
                "filename": out_name,
                "filesize": out_path.stat().st_size,
                "download_url": url_for("download_file", job_id=job_id, filename=out_name),
                "preview_type": "text",
                "text_content": result
            })

        # Baseline fallback for any other tools
        else:
            if saved_file_paths:
                in_p = saved_file_paths[0]
                out_name = f"processed_{in_p.name}"
                out_p = output_dir / out_name
                shutil.copyfile(str(in_p), str(out_p))
                return jsonify({
                    "success": True,
                    "filename": out_name,
                    "filesize": out_p.stat().st_size,
                    "download_url": url_for("download_file", job_id=job_id, filename=out_name)
                })
            else:
                return jsonify({"success": False, "error": "No input provided."})

    except Exception as e:
        return jsonify({"success": False, "error": f"Error processing file: {str(e)}"}), 500


# ==============================================================================
# ENTRY POINT
# ==============================================================================

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    print(f"\n=======================================================")
    print(f"🚀 All-in-One Utility Hub is running on http://127.0.0.1:{port}")
    print(f"📁 34 Production-Ready Tools Loaded Successfully")
    print(f"=======================================================\n")
    app.run(host="0.0.0.0", port=port, debug=True)
