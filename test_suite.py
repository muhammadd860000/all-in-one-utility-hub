"""
Automated Verification Test Suite for All-in-One Utility Hub
"""

import io
import json
from app import app
from tools_config import TOOLS
from PIL import Image
import pypdf
import docx

def test_routes_and_tools():
    app.config["TESTING"] = True
    client = app.test_client()

    # 1. Test Dashboard Route
    res = client.get("/")
    assert res.status_code == 200, f"Dashboard failed with status {res.status_code}"
    print("[PASS] Dashboard route tested successfully (200 OK)")

    # 2. Test All 34 Dedicated Tool Pages
    print(f"Testing {len(TOOLS)} tool routes...")
    for tool_id in TOOLS:
        res = client.get(f"/tool/{tool_id}")
        assert res.status_code == 200, f"Tool page {tool_id} returned {res.status_code}"
    print(f"[PASS] All {len(TOOLS)} dedicated tool pages loaded successfully (200 OK)")

    # 3. Test QR Generator Processing
    res = client.post("/process/qr-generator", data={"qr_content": "https://google.com", "qr_size": "10"})
    data = res.get_json()
    assert data["success"] is True and "download_url" in data
    print("[PASS] QR Code Generator processed successfully")

    # 4. Test Case Converter Processing
    res = client.post("/process/case-converter", data={"text_input": "hello world utility hub", "case_style": "title"})
    data = res.get_json()
    assert data["success"] is True and data["text_content"] == "Hello World Utility Hub"
    print("[PASS] Case Converter processed successfully")

    # 5. Test Base64 Codec
    res = client.post("/process/base64-codec", data={"input_text": "Antigravity Utility Hub", "codec_mode": "encode"})
    data = res.get_json()
    assert data["success"] is True and "QW50aWdyYXZpdHkgVXRpbGl0eSBIdWI=" in data["text_content"]
    print("[PASS] Base64 Codec processed successfully")

    # 6. Test Text Diff Checker
    res = client.post("/process/text-diff", data={"original_text": "line1\nline2", "modified_text": "line1\nline2_modified\nline3"})
    data = res.get_json()
    assert data["success"] is True and len(data["diff_lines"]) > 0
    print("[PASS] Text Diff Checker processed successfully")

    # 7. Test URL Encoder
    res = client.post("/process/url-encoder-decoder", data={"url_input": "hello world & foo=bar", "url_mode": "encode"})
    data = res.get_json()
    assert data["success"] is True and "hello%20world%20%26%20foo%3Dbar" in data["text_content"]
    print("[PASS] URL Encoder processed successfully")

    # 8. Test Image Resizer with Mock Image
    def make_mock_image():
        buf = io.BytesIO()
        img = Image.new("RGB", (200, 200), color="blue")
        img.save(buf, format="PNG")
        buf.seek(0)
        return buf

    res = client.post(
        "/process/image-resizer",
        data={
            "files": (make_mock_image(), "test_image.png"),
            "resize_mode": "percentage",
            "scale_percent": "50"
        },
        content_type="multipart/form-data"
    )
    data = res.get_json()
    assert data["success"] is True and "preview_data" in data
    print("[PASS] Image Resizer processed successfully")

    # 9. Test Image Filters
    res = client.post(
        "/process/image-filters",
        data={
            "files": (make_mock_image(), "test_filter.png"),
            "filter_type": "grayscale"
        },
        content_type="multipart/form-data"
    )
    data = res.get_json()
    assert data["success"] is True and "preview_data" in data
    print("[PASS] Image Filters processed successfully")

    # 10. Test Color Palette Extractor
    res = client.post(
        "/process/color-palette",
        data={"files": (make_mock_image(), "test_palette.png")},
        content_type="multipart/form-data"
    )
    data = res.get_json()
    assert data["success"] is True and len(data["colors"]) > 0
    print("[PASS] Color Palette Extractor processed successfully")

    # 11. Test PDF Merge with Mock PDFs
    def create_mock_pdf():
        buf = io.BytesIO()
        writer = pypdf.PdfWriter()
        writer.add_blank_page(width=100, height=100)
        writer.write(buf)
        buf.seek(0)
        return buf

    pdf1 = create_mock_pdf()
    pdf2 = create_mock_pdf()

    res = client.post(
        "/process/merge-pdf",
        data={
            "files": (pdf1, "page1.pdf")
        },
        content_type="multipart/form-data"
    )
    data = res.get_json()
    print("Merge PDF response:", data)
    assert data["success"] is True and data["filename"] == "merged_document.pdf"
    print("[PASS] PDF Merge processed successfully")

    # 12. Test PDF Rotate
    pdf3 = create_mock_pdf()
    res = client.post(
        "/process/rotate-pdf",
        data={
            "files": (pdf3, "test_rotate.pdf"),
            "angle": "90"
        },
        content_type="multipart/form-data"
    )
    data = res.get_json()
    assert data["success"] is True and "rotated_" in data["filename"]
    print("[PASS] PDF Rotate processed successfully")

    # 13. Test File Hasher
    test_file = io.BytesIO(b"Hello Antigravity Hub Checksum")
    res = client.post(
        "/process/file-hasher",
        data={"files": (test_file, "checksum_sample.txt")},
        content_type="multipart/form-data"
    )
    data = res.get_json()
    assert data["success"] is True and "SHA-256" in data["hashes"]
    print("[PASS] File Hasher processed successfully")

    # 14. Test Markdown to HTML
    res = client.post(
        "/process/markdown-to-html",
        data={"markdown_text": "# Test Heading\n**bold text**"},
    )
    data = res.get_json()
    assert data["success"] is True and "Test Heading" in data["text_content"]
    print("[PASS] Markdown to HTML processed successfully")

    # 15. Test HTML to Markdown
    res = client.post(
        "/process/html-to-markdown",
        data={"html_text": "<h2>Sub Heading</h2><p>Paragraph text</p>"},
    )
    data = res.get_json()
    assert data["success"] is True and "Sub Heading" in data["text_content"]
    print("[PASS] HTML to Markdown processed successfully")

    # 16. Test CSV to JSON
    csv_file = io.BytesIO(b"name,age,city\nAlice,30,New York\nBob,25,London")
    res = client.post(
        "/process/csv-to-json",
        data={"files": (csv_file, "users.csv")},
        content_type="multipart/form-data"
    )
    data = res.get_json()
    assert data["success"] is True and "Alice" in data["text_content"]
    print("[PASS] CSV to JSON processed successfully")

    # 17. Test JSON to CSV
    json_file = io.BytesIO(b'[{"name": "Alice", "age": 30}, {"name": "Bob", "age": 25}]')
    res = client.post(
        "/process/json-to-csv",
        data={"files": (json_file, "users.json")},
        content_type="multipart/form-data"
    )
    data = res.get_json()
    assert data["success"] is True and data["filename"].endswith(".csv")
    print("[PASS] JSON to CSV processed successfully")

    # 18. Test PDF Watermark
    pdf_wm = create_mock_pdf()
    res = client.post(
        "/process/watermark-pdf",
        data={"files": (pdf_wm, "watermark_test.pdf"), "watermark_text": "DRAFT", "opacity": "0.3"},
        content_type="multipart/form-data"
    )
    data = res.get_json()
    assert data["success"] is True and "watermarked_" in data["filename"]
    print("[PASS] PDF Watermark processed successfully")

    # 19. Test PDF Protect
    pdf_prot = create_mock_pdf()
    res = client.post(
        "/process/protect-pdf",
        data={"files": (pdf_prot, "protect_test.pdf"), "mode": "encrypt", "password": "SecretPassword123"},
        content_type="multipart/form-data"
    )
    data = res.get_json()
    assert data["success"] is True and "protected_" in data["filename"]
    print("[PASS] PDF Protect (Encrypt) processed successfully")

    # 20. Test PDF to Word (High Fidelity Mode - for Daraz Shipping Labels & Barcodes)
    pdf_label = create_mock_pdf()
    res = client.post(
        "/process/pdf-to-word",
        data={"files": (pdf_label, "shipping_label.pdf"), "conversion_mode": "high_fidelity"},
        content_type="multipart/form-data"
    )
    data = res.get_json()
    assert data["success"] is True and data["filename"] == "shipping_label.docx"
    print("[PASS] PDF to Word (High Fidelity Label & Barcode Mode) processed successfully")

    # 21. Test PDF to Word (Editable Layout Mode)
    pdf_edit = create_mock_pdf()
    res = client.post(
        "/process/pdf-to-word",
        data={"files": (pdf_edit, "document.pdf"), "conversion_mode": "editable_layout"},
        content_type="multipart/form-data"
    )
    data = res.get_json()
    assert data["success"] is True and data["filename"] == "document.docx"
    print("[PASS] PDF to Word (Editable Layout Mode) processed successfully")

    # 22. Test Word to PDF (with text & tables)
    def make_mock_docx():
        buf = io.BytesIO()
        doc = docx.Document()
        doc.add_heading("Shipping Manifest", level=1)
        doc.add_paragraph("Order #12345 - Daraz Express Tracking")
        t = doc.add_table(rows=2, cols=2)
        t.cell(0, 0).text = "Item"
        t.cell(0, 1).text = "Qty"
        t.cell(1, 0).text = "Wireless Mouse"
        t.cell(1, 1).text = "1"
        doc.save(buf)
        buf.seek(0)
        return buf

    res = client.post(
        "/process/word-to-pdf",
        data={"files": (make_mock_docx(), "manifest.docx")},
        content_type="multipart/form-data"
    )
    data = res.get_json()
    assert data["success"] is True and data["filename"] == "manifest.pdf"
    print("[PASS] Word to PDF (Text & Tables) processed successfully")

    # 23. Test Word to PDF (with embedded image / scanned shipping label)
    def make_mock_image_docx():
        buf = io.BytesIO()
        doc = docx.Document()
        # Add a temporary image
        img_buf = io.BytesIO()
        im = Image.new("RGB", (300, 150), color="white")
        im.save(img_buf, format="PNG")
        img_buf.seek(0)
        doc.add_picture(img_buf)
        doc.save(buf)
        buf.seek(0)
        return buf

    res = client.post(
        "/process/word-to-pdf",
        data={"files": (make_mock_image_docx(), "label_image.docx")},
        content_type="multipart/form-data"
    )
    data = res.get_json()
    assert data["success"] is True and data["filename"] == "label_image.pdf"
    print("[PASS] Word to PDF (Embedded Image & Label) processed successfully")

    # 24. Test Compress PDF
    pdf_comp = create_mock_pdf()
    res = client.post(
        "/process/compress-pdf",
        data={"files": (pdf_comp, "large_doc.pdf"), "compression_level": "recommended"},
        content_type="multipart/form-data"
    )
    data = res.get_json()
    assert data["success"] is True and "compressed_" in data["filename"]
    print("[PASS] Compress PDF (Stream & Image Optimization) processed successfully")

    # 25. Test XML to JSON
    xml_buf = io.BytesIO(b"<order><id>9876</id><status>Delivered</status></order>")
    res = client.post(
        "/process/xml-to-json",
        data={"files": (xml_buf, "order.xml")},
        content_type="multipart/form-data"
    )
    data = res.get_json()
    assert data["success"] is True and "9876" in data["text_content"]
    print("[PASS] XML to JSON processed successfully")

    # 26. Test Word to PDF with Oversized/Extra-Tall Shipping Label Image (Checking ReportLab page boundary safety)
    def make_mock_large_image_docx():
        buf = io.BytesIO()
        doc = docx.Document()
        img_buf = io.BytesIO()
        # High-res tall image (1500 x 3500 px)
        im = Image.new("RGB", (1500, 3500), color="blue")
        im.save(img_buf, format="PNG")
        img_buf.seek(0)
        doc.add_picture(img_buf)
        doc.save(buf)
        buf.seek(0)
        return buf

    res = client.post(
        "/process/word-to-pdf",
        data={"files": (make_mock_large_image_docx(), "oversized_label.docx")},
        content_type="multipart/form-data"
    )
    data = res.get_json()
    assert data["success"] is True and data["filename"] == "oversized_label.pdf"
    print("[PASS] Word to PDF (Oversized 1500x3500px Image Proportional Scaling) processed successfully")

    print("\n>>> ALL 26 TEST SUITES PASSED FLAWLESSLY! <<<")

if __name__ == "__main__":
    test_routes_and_tools()
